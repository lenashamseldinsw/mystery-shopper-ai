from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re
import os


class DOCXBuilder:
    def __init__(self):
        self.font = "Arial"

    # -------------------------------------------------
    # Base document
    # -------------------------------------------------
    def create_document(self):
        doc = Document()
        section = doc.sections[0]
        section.right_margin = Inches(1)
        section.left_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        return doc

    # -------------------------------------------------
    # Logo
    # -------------------------------------------------
    def add_logo_header(self, doc, logo_path):
        if logo_path and os.path.exists(logo_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(logo_path, width=Inches(2.5))

    # -------------------------------------------------
    # Titles & headers (NOT lists)
    # -------------------------------------------------
    def add_title_section(self, doc, text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.right_to_left = True

        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = self.font

    def add_section_header(self, doc, text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_to_left = True

        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = self.font

    # -------------------------------------------------
    # Footer
    # -------------------------------------------------
    def add_footer(self, doc, text):
        footer = doc.sections[0].footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.right_to_left = True

        run = p.add_run(text)
        run.font.name = self.font
        run.font.size = Pt(9)

    # -------------------------------------------------
    # LISTS — Word-native ONLY (this is the fix)
    # -------------------------------------------------

    # ✅ This is the SAME bullet behavior from the point where bullets worked
    def _add_bullet_item(self, doc, text):
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_to_left = True

        run = p.add_run(text)
        run.font.name = self.font
        run.font.size = Pt(12)

    # ✅ Fixed numbering: Word-native, RTL, no heading abuse
    def _add_numbered_item(self, doc, text):
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_to_left = True

        run = p.add_run(text)
        run.font.name = self.font
        run.font.size = Pt(12)

    # -------------------------------------------------
    # Dispatcher
    # -------------------------------------------------
    def add_formatted_paragraph(self, doc, content):
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph("")
                continue

            # Remove ALL markdown formatting (stronger removal)
            line = re.sub(r"\*\*(.*?)\*\*", r"\1", line)  # **bold**
            line = re.sub(r"\*(.*?)\*", r"\1", line)      # *italic*
            line = re.sub(r"\*+", "", line)               # Remove any remaining asterisks
            
            # Remove hashtags from LLM text
            line = re.sub(r"#{1,6}\s*", "", line)         # Remove ### headings
            
            # Fix percentage formatting: 70۔0% → 70%, %0-70 → 70%
            line = re.sub(r"(\d+)۔(\d+)%", r"\1.\2%", line)  # 70۔0% → 70.0%
            line = re.sub(r"(\d+)\.0%", r"\1%", line)        # 70.0% → 70%
            line = re.sub(r"%(\d+)-(\d+)", r"\2%", line)     # %0-70 → 70%
            line = re.sub(r"%(\d+)", r"\1%", line)           # %70 → 70%
            
            # Section headers (###)
            if line.startswith("###"):
                self.add_section_header(
                    doc,
                    line.replace("###", "").strip()
                )

            # Numbered list: 1. text -> Convert to HEADING instead of numbered list
            elif re.match(r"^\d+\.\s+", line):
                # Remove number and make it a heading
                heading_text = re.sub(r"^\d+\.\s+", "", line).strip()
                if heading_text:
                    self.add_section_header(doc, heading_text)

            # Bullet list: • text / * text / - text -> Use RTL-configured bullet
            elif re.match(r"^[•*\-●][\s\t]+", line):
                # Remove bullet character (including tabs) and create RTL bullet
                bullet_text = re.sub(r"^[•*\-●][\s\t]+", "", line).strip()
                # Also remove any remaining asterisks from LLM formatting
                bullet_text = re.sub(r"\*+", "", bullet_text)
                
                if bullet_text:  # Only process if there's text after removing bullet
                    # Check if text contains a heading pattern BEFORE period replacement
                    # Common patterns: "text:" or just bold keywords
                    heading_match = re.match(r'^([^:：]+[:：])\s*(.*)$', bullet_text)
                    
                    # Replace periods with Arabic periods and fix percentages
                    if heading_match:
                        heading_part = heading_match.group(1).strip().replace(".", "۔")
                        body_part = heading_match.group(2).strip().replace(".", "۔")
                        # Fix percentages
                        heading_part = re.sub(r"(\d+)۔(\d+)%", r"\1.\2%", heading_part)
                        heading_part = re.sub(r"(\d+)\.0%", r"\1%", heading_part)
                        body_part = re.sub(r"(\d+)۔(\d+)%", r"\1.\2%", body_part)
                        body_part = re.sub(r"(\d+)\.0%", r"\1%", body_part)
                        # Remove any trailing dashes or em dashes
                        heading_part = re.sub(r'[\s\-—۔]+$', '', heading_part)
                        body_part = re.sub(r'[\s\-—۔]+$', '', body_part)
                    else:
                        # No heading pattern found
                        heading_part = None
                        body_part = bullet_text.replace(".", "۔")
                        # Fix percentages
                        body_part = re.sub(r"(\d+)۔(\d+)%", r"\1.\2%", body_part)
                        body_part = re.sub(r"(\d+)\.0%", r"\1%", body_part)
                        # Remove any trailing dashes or em dashes
                        body_part = re.sub(r'[\s\-—۔]+$', '', body_part)
                    
                    # Create paragraph with bullet style
                    p = doc.add_paragraph(style='List Bullet')
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Right alignment for Arabic
                    p.paragraph_format.right_to_left = True  # RTL paragraph
                    
                    # Get paragraph properties XML element
                    pPr = p._element.get_or_add_pPr()
                    
                    # Add explicit RTL direction to paragraph
                    bidi = OxmlElement('w:bidi')
                    bidi.set(qn('w:val'), '1')
                    pPr.append(bidi)
                    
                    # Set right justification at XML level
                    jc = OxmlElement('w:jc')
                    jc.set(qn('w:val'), 'right')
                    pPr.append(jc)
                    
                    # Set RTL-specific indentation for bullets
                    # Remove default indent first
                    ind = pPr.find(qn('w:ind'))
                    if ind is not None:
                        pPr.remove(ind)
                    
                    # Add new indent with right alignment
                    ind = OxmlElement('w:ind')
                    ind.set(qn('w:right'), '0')        # No extra right margin - text starts from right
                    ind.set(qn('w:hanging'), '0')      # No hanging indent
                    pPr.append(ind)
                    
                    # Add content with proper formatting
                    if heading_part:
                        # Add heading part (bold and italic)
                        heading_run = p.add_run(heading_part)
                        heading_run.font.name = self.font
                        heading_run.font.size = Pt(12)
                        heading_run.bold = True
                        heading_run.italic = True
                        
                        # Add space and body part (normal)
                        if body_part:
                            space_run = p.add_run(" ")
                            space_run.font.name = self.font
                            space_run.font.size = Pt(12)
                            
                            body_run = p.add_run(body_part)
                            body_run.font.name = self.font
                            body_run.font.size = Pt(12)
                    else:
                        # No heading, just add the text normally
                        run = p.add_run(body_part)
                        run.font.name = self.font
                        run.font.size = Pt(12)

            # Normal paragraph
            else:
                # Remove any remaining asterisks from LLM formatting
                clean_line = re.sub(r"\*+", "", line)
                
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p.paragraph_format.right_to_left = True

                # Replace periods with Arabic periods (but preserve percentages)
                text_with_arabic_periods = clean_line.replace(".", "۔")
                # Fix percentage formatting that might have been affected
                text_with_arabic_periods = re.sub(r"(\d+)۔(\d+)%", r"\1.\2%", text_with_arabic_periods)
                text_with_arabic_periods = re.sub(r"(\d+)\.0%", r"\1%", text_with_arabic_periods)
                
                run = p.add_run(text_with_arabic_periods)
                run.font.name = self.font
                run.font.size = Pt(12)

    # -------------------------------------------------
    # Save
    # -------------------------------------------------
    def save_to_buffer(self, doc):
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf


class ContentProcessor:
    @staticmethod
    def parse_sections_from_text(text):
        sections = {}
        current = None
        buff = []

        for line in text.split("\n"):
            if line.startswith("==="):
                if current:
                    sections[current] = "\n".join(buff).strip()
                current = line.replace("=", "").strip()
                buff = []
            else:
                buff.append(line)

        if current:
            sections[current] = "\n".join(buff).strip()

        return sections
